"""
Gmail OCR Extraction Pipeline – Enhanced Business Email Tracking
Saves attachments to Drive + Background processing with comprehensive email filtering.
Stores email text + locally extracted attachment content in MySQL.
Enhanced OCR and text extraction for all attachment types.

Fixes applied (v2):
  1. Removed hardcoded credentials – all config MUST come from config.ini.
  2. find_file_in_folder: removed misleading file_name param; only hash-based lookup.
  3. setup_tesseract() result is cached via module-level flag (_TESSERACT_AVAILABLE).
  4. Large file_bytes no longer sent over ProcessPoolExecutor IPC – written to a
     temp file whose path is passed instead, avoiding large pickle copies.
  5. get_executor() registers an atexit handler to shut down the pool cleanly.
  6. upsert_thread: column names are validated against an allowlist to prevent
     accidental SQL breakage.
  7. load_existing_rows: accepts an optional set of thread_ids to load selectively.
  8. run_continuous: resumes by thread_id stored in DB, not by numeric index.
  9. thread_has_changed_since: passes historyTypes filter to the History API.
 10. parse_email_date / top-level imports cleaned up.
 11. OCR_MAX_CHARS / BODY_MAX_CHARS moved to module-level constants.
 12. CC deduplication now extracts bare email addresses before deduping.
 13. init_db ALTER TABLE list now includes drive_folder_id and contacts/footprint cols.
 14. generate_report accepts and returns data cleanly; caller handles printing.
 15. Worker-process imports are guarded so missing optional deps fail gracefully.
"""

import atexit
import base64
import io
import csv
import concurrent.futures
import json
import logging
import os
import re
import sys
import time
import hashlib
import tempfile
import subprocess
import traceback
from datetime import datetime, timedelta, timezone
from email.utils import parsedate_to_datetime          # FIX 10: top-level import
from typing import Dict, Iterable, List, Optional, Set, Tuple

import mysql.connector
from google.auth.transport.requests import Request
from google.oauth2.credentials import Credentials
from google_auth_oauthlib.flow import InstalledAppFlow
from googleapiclient.discovery import build
from googleapiclient.errors import HttpError
from googleapiclient.http import MediaIoBaseUpload

# -----------------------------------------------------------------------
# MODULE-LEVEL CONSTANTS  (FIX 11)
# -----------------------------------------------------------------------
BODY_MAX_CHARS = 100_000
OCR_MAX_CHARS  = 200_000

# -----------------------------------------------------------------------
# PROCESS POOL  (FIX 4 + 5)
# -----------------------------------------------------------------------
EXECUTOR: Optional[concurrent.futures.ProcessPoolExecutor] = None

def get_executor() -> concurrent.futures.ProcessPoolExecutor:
    global EXECUTOR
    if EXECUTOR is None:
        EXECUTOR = concurrent.futures.ProcessPoolExecutor(max_workers=4)
        atexit.register(_shutdown_executor)          # FIX 5: clean shutdown
    return EXECUTOR

def _shutdown_executor():
    global EXECUTOR
    if EXECUTOR is not None:
        EXECUTOR.shutdown(wait=False)
        EXECUTOR = None

# -----------------------------------------------------------------------
# CONFIGURATION  (FIX 1: no hardcoded credentials)
# -----------------------------------------------------------------------
import configparser

_cfg = configparser.ConfigParser()
_cfg.read("config.ini")

def _require(section: str, key: str) -> str:
    """Read a required config value; raise clearly if missing."""
    try:
        return _cfg.get(section, key)
    except (configparser.NoSectionError, configparser.NoOptionError):
        raise RuntimeError(
            f"Missing required config value [{section}] {key} in config.ini"
        )

MYSQL_HOST     = _cfg.get("mysql", "host",     fallback=None) or _require("mysql", "host")
MYSQL_USER     = _cfg.get("mysql", "user",     fallback=None) or _require("mysql", "user")
MYSQL_PASSWORD = _cfg.get("mysql", "password", fallback=None) or _require("mysql", "password")
MYSQL_DATABASE = _cfg.get("mysql", "database", fallback=None) or _require("mysql", "database")
MYSQL_PORT     = _cfg.getint("mysql", "port",  fallback=3306)

DRIVE_FOLDER_ID = _cfg.get("drive", "folder_id", fallback=None) or _require("drive", "folder_id")

print(f"Using MySQL host={MYSQL_HOST}, user={MYSQL_USER}, database={MYSQL_DATABASE}, port={MYSQL_PORT}")

SCOPES = [
    "https://www.googleapis.com/auth/gmail.readonly",
    "https://www.googleapis.com/auth/drive.file",
    "https://www.googleapis.com/auth/documents.readonly",
    "https://www.googleapis.com/auth/spreadsheets.readonly",
    "https://www.googleapis.com/auth/presentations.readonly",
]

IMPORTANT_DOMAINS = [
    'gov.in', 'nic.in', 'gem.gov.in', 'eprocure.gov.in', 'cppp.gov.in',
    'bhel.com', 'ntpc.co.in', 'ongcindia.com', 'iocl.com', 'bpcl.co.in',
    'hpcl.co.in', 'gail.co.in', 'sail.co.in', 'coalindia.in', 'nhpc.nic.in',
    'larsentoubro.com', 'tata.com', 'relianceada.com', 'adanienterprises.com',
    'mahindra.com', 'godrej.com', 'birla.com', 'jindalsteel.com',
    'hdfcbank.com', 'icicibank.com', 'sbi.co.in', 'axisbank.com',
    'pnb.co.in', 'bankofbaroda.in', 'canarabank.com',
]



MAX_RUN_TIME_MS      = 25 * 60 * 1000
MAX_THREADS_PER_RUN  = 500
BATCH_SIZE           = 500
LOOKBACK_DAYS        = 1

_SEARCH_BASE = (
    "-in:trash "
    "-from:corporatenetbanking.automailer@hdfcbank.bank.in"
)
_SEARCH_UNTIL = datetime.now().strftime("%Y/%m/%d")
_SEARCH_SINCE = (datetime.now() - timedelta(days=LOOKBACK_DAYS)).strftime("%Y/%m/%d")

SEARCH_QUERY = f"{_SEARCH_BASE} after:{_SEARCH_SINCE} before:{_SEARCH_UNTIL}"

# -----------------------------------------------------------------------
# LOGGING
# -----------------------------------------------------------------------
logging.basicConfig(
    level=logging.INFO,
    format="%(asctime)s [%(levelname)s] %(message)s",
    handlers=[
        logging.StreamHandler(sys.stdout),
        logging.FileHandler('gmail_pipeline.log', encoding='utf-8'),
    ],
)
logger = logging.getLogger(__name__)

# -----------------------------------------------------------------------
# ALLOWED DB COLUMNS  (FIX 6: prevent accidental bad SQL)
# -----------------------------------------------------------------------
_ALLOWED_THREAD_COLUMNS: Set[str] = {
    "thread_id", "related_ids", "msg_count", "date", "sender",
    "sender_details", "cc_details", "subject", "body", "attach_names",
    "attach_links", "ocr_text", "ai_summary", "category", "sub_category",
    "priority", "is_important", "importance_reasons", "contacts",
    "footprint", "message_ids", "history_ids", "latest_history_id",
    "drive_folder_id", "company", "to_details",
}

# -----------------------------------------------------------------------
# DATABASE SETUP
# -----------------------------------------------------------------------
def get_db_connection():
    return mysql.connector.connect(
        host=MYSQL_HOST,
        user=MYSQL_USER,
        password=MYSQL_PASSWORD,
        database=MYSQL_DATABASE,
        port=MYSQL_PORT,
        charset='utf8mb4',
        use_unicode=True,
        autocommit=False,
    )

def init_db():
    conn   = get_db_connection()
    cursor = conn.cursor()

    try:
        cursor.execute("""
            CREATE TABLE IF NOT EXISTS threads (
                id                  INT AUTO_INCREMENT PRIMARY KEY,
                thread_id           VARCHAR(255) NOT NULL UNIQUE,
                related_ids         TEXT,
                msg_count           INT DEFAULT 0,
                date                TIMESTAMP NULL,
                sender              VARCHAR(512),
                sender_details      VARCHAR(512),
                cc_details          TEXT,
                subject             TEXT,
                body                LONGTEXT,
                attach_names        TEXT,
                attach_links        TEXT,
                ocr_text            LONGTEXT,
                ai_summary          LONGTEXT,
                category            VARCHAR(100),
                sub_category        VARCHAR(100),
                priority            VARCHAR(20),
                is_important        BOOLEAN DEFAULT FALSE,
                importance_reasons  TEXT,
                contacts            TEXT,
                footprint           LONGTEXT,
                last_updated        TIMESTAMP DEFAULT CURRENT_TIMESTAMP ON UPDATE CURRENT_TIMESTAMP,
                message_ids         TEXT,
                history_ids         TEXT,
                latest_history_id   VARCHAR(255),
                drive_folder_id     VARCHAR(255),
                INDEX idx_thread_id       (thread_id),
                INDEX idx_date            (date),
                INDEX idx_category        (category),
                INDEX idx_is_important    (is_important),
                INDEX idx_priority        (priority)
            ) CHARACTER SET utf8mb4 COLLATE utf8mb4_unicode_ci;
        """)

        # FIX 13: drive_folder_id and other missing columns added to ALTER list
        new_columns = [
            ("sub_category",        "VARCHAR(100)"),
            ("priority",            "VARCHAR(20)"),
            ("is_important",        "BOOLEAN DEFAULT FALSE"),
            ("importance_reasons",  "TEXT"),
            ("history_ids",         "TEXT"),
            ("latest_history_id",   "VARCHAR(255)"),
            ("drive_folder_id",     "VARCHAR(255)"),
            ("contacts",            "TEXT"),
            ("footprint",           "LONGTEXT"),
            ("company",             "VARCHAR(500) DEFAULT NULL"),
            ("to_details",          "TEXT"),
        ]
        for col_name, col_type in new_columns:
            try:
                cursor.execute(f"ALTER TABLE threads ADD COLUMN {col_name} {col_type}")
                logger.info(f"Added column {col_name} to threads table.")
            except mysql.connector.errors.ProgrammingError:
                pass  # already exists

        # Ensure to_details is TEXT (idempotent upgrade from any prior VARCHAR)
        try:
            cursor.execute("ALTER TABLE threads MODIFY to_details TEXT")
            logger.info("Upgraded to_details to TEXT")
        except mysql.connector.errors.ProgrammingError:
            pass

        cursor.execute("""
            CREATE TABLE IF NOT EXISTS processing_status (
                id                  INT AUTO_INCREMENT PRIMARY KEY,
                total_threads       INT DEFAULT 0,
                processed_threads   INT DEFAULT 0,
                important_threads   INT DEFAULT 0,
                last_thread_id      VARCHAR(255),
                last_processed_date TIMESTAMP NULL,
                batch_number        INT DEFAULT 0,
                start_time          TIMESTAMP NULL,
                end_time            TIMESTAMP NULL,
                status              VARCHAR(50) DEFAULT 'running',
                created_at          TIMESTAMP DEFAULT CURRENT_TIMESTAMP
            );
        """)

        cursor.execute("""
            CREATE TABLE IF NOT EXISTS email_categories (
                id            INT AUTO_INCREMENT PRIMARY KEY,
                category_name VARCHAR(100) UNIQUE,
                keywords      TEXT,
                priority      VARCHAR(20) DEFAULT 'medium',
                created_at    TIMESTAMP DEFAULT CURRENT_TIMESTAMP
            );
        """)

        cursor.execute("SELECT COUNT(*) FROM email_categories")
        if cursor.fetchone()[0] == 0:
            default_categories = [
                ('Tender/RFP/Bid',      'tender,rfp,rfq,eoi,bid,procurement,gem,notice inviting,expression of interest,request for proposal,request for quotation,corrigendum,pre-bid,technical bid,commercial bid,financial bid,price bid,earnest money,emd,security deposit,performance guarantee,limited tender,open tender,single tender,two bid,nit,tender id,tender no,nit no,bidder,submission of bid,reverse auction,ra,clarification', 'high'),
                ('Purchase Order',      'purchase order,po,po#,po number,work order,wo,supply order,letter of award,loa,letter of intent,loi,rate contract', 'high'),
                ('Contract/Agreement',  'contract,agreement,mou,nda,non-disclosure,memorandum of understanding,partnership,joint venture,collaboration,service level agreement,sla', 'high'),
                ('Invoice/Billing',     'invoice,billing,bill,proforma,payment advice,debit note,credit note,tax invoice,commercial invoice,outstanding,receivable', 'medium'),
                ('Banking/Finance',     'bank statement,bank guarantee,bg,letter of credit,lc,financial,emd,security deposit,pbg,performance guarantee,bank confirmation,neft,rtgs,imps,remittance,transaction,disbursement,disbursed,bank account,fixed deposit,fdr,dd,demand draft', 'high'),
                ('Client Communication','client,customer,project update,delivery schedule,meeting,requirement,feedback,review,approval,minutes of meeting,mom,site visit', 'medium'),
                ('Vendor/Supplier',     'vendor,supplier,material,equipment,quotation,rate,price list,catalogue,brochure,specifications,oem,distributor', 'medium'),
                ('Legal/Compliance',    'legal,compliance,regulatory,tax,gst,income tax,audit,statutory,license,registration,certification,iso,quality,arbitration,litigation', 'high'),
                ('Government/PSU',      'government,ministry,department,psu,public sector,undertaking,corporation,board,commission,authority,gov.in,nic.in', 'high'),
                ('General',             '', 'low'),
            ]
            for cat_name, keywords, priority in default_categories:
                cursor.execute(
                    "INSERT INTO email_categories (category_name, keywords, priority) VALUES (%s, %s, %s)",
                    (cat_name, keywords, priority),
                )
            logger.info("Inserted default email categories.")

        conn.commit()
        logger.info("Database tables are ready.")
    except mysql.connector.Error as err:
        if getattr(err, 'errno', None) == 1290 or 'read-only' in str(err).lower():
            logger.warning(
                "Database is in read-only mode; skipping schema initialization and data inserts. "
                "Make sure required tables already exist or run this against a writable instance."
            )
            conn.rollback()
        else:
            raise
    finally:
        cursor.close()
        conn.close()

# -----------------------------------------------------------------------
# GOOGLE API AUTH
# -----------------------------------------------------------------------
def safe_remove_path(path):
    try:
        if os.path.isdir(path):
            import shutil
            shutil.rmtree(path)
        elif os.path.exists(path):
            os.remove(path)
    except Exception as e:
        logger.warning(f"Could not remove path {path}: {e}")

def find_root_auth_file(filename):
    candidates = [
        os.path.join("root_config", filename),
        os.path.join("/app", "root_config", filename),
        filename,
        os.path.join("..", filename),
        os.path.join(os.path.dirname(__file__), filename),
        os.path.join(os.path.dirname(__file__), "..", filename)
    ]
    for c in candidates:
        if os.path.isfile(c):
            return c
    return filename

def get_google_services():
    creds = None
    token_file = find_root_auth_file("token.json")
    creds_file = find_root_auth_file("credentials.json")

    if os.path.isdir(token_file):
        safe_remove_path(token_file)
    if os.path.isdir(creds_file):
        safe_remove_path(creds_file)

    if os.path.isfile(token_file):
        try:
            creds = Credentials.from_authorized_user_file(token_file, SCOPES)
            if creds and creds.valid:
                if not set(SCOPES).issubset(set(creds.scopes)):
                    logger.warning("Token scopes outdated. Re-authenticating...")
                    safe_remove_path(token_file)
                    creds = None
        except Exception:
            logger.warning("Error loading token. Re-authenticating...")
            safe_remove_path(token_file)
            creds = None

    if not creds or not creds.valid:
        if creds and creds.expired and creds.refresh_token:
            try:
                creds.refresh(Request())
            except Exception:
                logger.warning("Token refresh failed. Re-authenticating...")
                safe_remove_path(token_file)
                creds = None
        if not creds:
            if not os.path.isfile(creds_file):
                raise FileNotFoundError(f"{creds_file} is missing or not a valid file.")
            flow  = InstalledAppFlow.from_client_secrets_file(creds_file, SCOPES)
            creds = flow.run_local_server(port=0)
        with open(token_file, "w") as token:
            token.write(creds.to_json())
        logger.info("New token saved.")

    gmail_service = build("gmail", "v1", credentials=creds)
    drive_service = build("drive", "v3", credentials=creds)
    logger.info("Google services authenticated.")
    return gmail_service, drive_service

# -----------------------------------------------------------------------
# EMAIL CLASSIFICATION
# -----------------------------------------------------------------------
def determine_importance(
    subject: str, body: str, sender: str, attachments: List[str]
) -> Tuple[bool, List[str]]:
    subject_lower  = (subject or "").lower()
    body_lower     = (body    or "").lower()
    sender_lower   = (sender  or "").lower()
    importance_reasons: List[str] = []
    is_important   = False

    for domain in IMPORTANT_DOMAINS:
        if domain in sender_lower:
            importance_reasons.append(f"Important sender domain: {domain}")
            is_important = True
            break

    if attachments and any(
        a.lower().endswith(('.pdf', '.doc', '.docx', '.xls', '.xlsx'))
        for a in attachments
    ):
        importance_reasons.append("Contains important attachment")
        is_important = True

    category_keywords = {
        'Tender/RFP/Bid': {
            'keywords': ['tender','rfp','rfq','eoi','bid','procurement','gem',
                         'notice inviting','expression of interest','request for proposal',
                         'corrigendum','pre-bid','technical bid','commercial bid',
                         'financial bid','price bid','earnest money','emd',
                         'security deposit','performance guarantee','limited tender',
                         'open tender','single tender','two bid','nit'],
            'priority': 'high',
        },
        'Purchase Order': {
            'keywords': ['purchase order','po#','po number','work order',
                         'supply order','letter of award','loa','letter of intent',
                         'loi','rate contract'],
            'priority': 'high',
        },
        'Contract/Agreement': {
            'keywords': ['contract','agreement','mou','nda','non-disclosure',
                         'memorandum of understanding','partnership','joint venture',
                         'collaboration','service level agreement','sla'],
            'priority': 'high',
        },
        'Invoice/Billing': {
            'keywords': ['invoice','bill','proforma','payment','debit note',
                         'credit note','tax invoice','commercial invoice'],
            'priority': 'medium',
        },
        'Banking/Finance': {
            'keywords': ['bank statement','bank guarantee','bg','letter of credit',
                         'lc','financial','emd','security deposit',
                         'performance guarantee','bank confirmation'],
            'priority': 'high',
        },
        'Client Communication': {
            'keywords': ['client','customer','project update','delivery schedule',
                         'meeting','requirement','feedback','review','approval'],
            'priority': 'medium',
        },
        'Vendor/Supplier': {
            'keywords': ['vendor','supplier','material','equipment','quotation',
                         'rate','price list','catalogue','brochure','specifications'],
            'priority': 'medium',
        },
        'Legal/Compliance': {
            'keywords': ['legal','compliance','regulatory','tax','gst',
                         'income tax','audit','statutory','license','registration',
                         'certification','iso','quality'],
            'priority': 'high',
        },
        'Government/PSU': {
            'keywords': ['government','ministry','department','psu','public sector',
                         'undertaking','corporation','board','commission','authority',
                         'gov.in','nic.in'],
            'priority': 'high',
        },
        'Sales/Proposal': {
            'keywords': ['sales','proposal','offer','discount','negotiation',
                         'deal','opportunity','lead','prospect','enquiry'],
            'priority': 'medium',
        },
    }

    max_matches   = 0
    best_category = "General"
    best_priority = "low"
    for cat_name, cat_data in category_keywords.items():
        matches = sum(1 for kw in cat_data['keywords'] if kw in subject_lower or kw in body_lower)
        if matches > max_matches:
            max_matches   = matches
            best_category = cat_name
            best_priority = cat_data['priority']

    category = best_category
    if best_priority == 'high':
        is_important = True
        importance_reasons.append(f"High priority category: {category}")

    important_patterns = [
        (r'(?:deadline|due date|last date).*?(\d{1,2}[-/]\d{1,2}[-/]\d{2,4})', 'Contains deadline'),
        (r'(?:urgent|immediate|critical|priority|asap)',                          'Marked as urgent/priority'),
        (r'(?:confidential|private|sensitive)',                                   'Marked as confidential'),
        (r'(?:amount|value|cost|price).*?(?:rs\.?|inr|₹)\s*\d[\d,]*',           'Contains monetary value'),
        (r'(?:meeting|conference call|video call).*?(?:scheduled|arranged|fixed)','Meeting scheduled'),
    ]
    combined_text = f"{subject_lower} {body_lower[:5000]}"
    for pattern, reason in important_patterns:
        if re.search(pattern, combined_text, re.IGNORECASE):
            importance_reasons.append(reason)
            is_important = True

    return is_important, list(set(importance_reasons))

# -----------------------------------------------------------------------
# DRIVE HELPERS
# -----------------------------------------------------------------------
def create_drive_folder(drive_service, folder_name: str, parent_folder_id: str = None) -> Optional[str]:
    try:
        meta = {'name': folder_name, 'mimeType': 'application/vnd.google-apps.folder'}
        if parent_folder_id:
            meta['parents'] = [parent_folder_id]
        folder = drive_service.files().create(body=meta, fields='id').execute()
        logger.info(f"Created Drive folder: {folder_name}")
        return folder.get('id')
    except Exception as e:
        logger.error(f"Failed to create Drive folder: {e}")
        return None

def get_or_create_thread_folder(drive_service, thread_id: str, root_folder_id: str) -> str:
    folder_name = f"thread_{thread_id.replace('/', '_')}"
    try:
        query   = (
            f"name='{folder_name}' and "
            f"mimeType='application/vnd.google-apps.folder' and "
            f"'{root_folder_id}' in parents and trashed=false"
        )
        results = drive_service.files().list(q=query, fields='files(id)', pageSize=1).execute()
        if results.get('files'):
            return results['files'][0]['id']
        return create_drive_folder(drive_service, folder_name, root_folder_id) or root_folder_id
    except Exception as e:
        logger.error(f"Error getting/creating thread folder: {e}")
        return root_folder_id

def find_file_in_folder_by_hash(drive_service, folder_id: str, file_hash: str) -> Optional[dict]:
    """FIX 2: only hash-based lookup; file_name param removed (was unused)."""
    try:
        query   = (
            f"appProperties has {{ key='sha256' and value='{file_hash}' }} "
            f"and '{folder_id}' in parents and trashed=false"
        )
        results = drive_service.files().list(
            q=query, fields='files(id, name, webViewLink)', pageSize=1
        ).execute()
        files = results.get('files', [])
        return files[0] if files else None
    except Exception as e:
        logger.error(f"Failed to check for existing file by hash {file_hash[:8]}…: {e}")
        return None

def upload_to_drive_dedup(
    drive_service, file_bytes: bytes, file_name: str, parent_folder_id: str
) -> Optional[str]:
    """Upload only if not already present by content hash. Returns webViewLink."""
    file_hash = hashlib.sha256(file_bytes).hexdigest()
    existing  = find_file_in_folder_by_hash(drive_service, parent_folder_id, file_hash)
    if existing:
        logger.info(f"'{file_name}' already on Drive (hash match). Reusing link.")
        return existing.get('webViewLink')

    mime_types = {
        'pdf':  'application/pdf',
        'doc':  'application/msword',
        'docx': 'application/vnd.openxmlformats-officedocument.wordprocessingml.document',
        'xlsx': 'application/vnd.openxmlformats-officedocument.spreadsheetml.sheet',
        'xls':  'application/vnd.ms-excel',
        'pptx': 'application/vnd.openxmlformats-officedocument.presentationml.presentation',
        'ppt':  'application/vnd.ms-powerpoint',
        'jpg':  'image/jpeg', 'jpeg': 'image/jpeg',
        'png':  'image/png',  'gif':  'image/gif',
        'bmp':  'image/bmp',  'tiff': 'image/tiff', 'webp': 'image/webp',
        'txt':  'text/plain', 'csv':  'text/csv',   'rtf':  'application/rtf',
    }
    file_ext  = file_name.rsplit('.', 1)[-1].lower() if '.' in file_name else ''
    mime_type = mime_types.get(file_ext, 'application/octet-stream')

    try:
        file_metadata = {
            'name':          file_name,
            'parents':       [parent_folder_id],
            'appProperties': {'sha256': file_hash},
        }
        media = MediaIoBaseUpload(io.BytesIO(file_bytes), mimetype=mime_type, resumable=True)
        file  = drive_service.files().create(
            body=file_metadata, media_body=media, fields='id, webViewLink'
        ).execute()
        logger.info(f"Uploaded '{file_name}' to Drive.")
        return file.get('webViewLink')
    except Exception as e:
        logger.error(f"Failed to upload {file_name}: {e}")
        return None

# -----------------------------------------------------------------------
# GMAIL HELPERS
# -----------------------------------------------------------------------
def search_threads(gmail_service, query: str) -> List[dict]:
    threads    = []
    page_token = None
    page_count = 0
    retry_count = 0
    while True:
        try:
            response = gmail_service.users().threads().list(
                userId="me", q=query, pageToken=page_token, maxResults=500
            ).execute()
            page = response.get("threads", [])
            threads.extend(page)
            page_token = response.get("nextPageToken")
            page_count += 1
            logger.info(f"  search_threads: page {page_count}, {len(threads)} threads so far")
            if not page_token:
                break
            retry_count = 0
        except HttpError as e:
            if e.resp.status in [429, 500, 503] and retry_count < 5:
                wait = (2 ** retry_count) * 2
                logger.warning(f"API rate limit/error, retrying in {wait}s…")
                time.sleep(wait)
                retry_count += 1
            else:
                logger.error(f"Failed to search threads: {e}")
                break
    return threads

def get_thread_messages(gmail_service, thread_id: str) -> List[dict]:
    try:
        thread = gmail_service.users().threads().get(userId="me", id=thread_id).execute()
        return thread.get("messages", [])
    except Exception as e:
        logger.error(f"Failed to get thread {thread_id}: {e}")
        return []

def get_message_data(gmail_service, message_id: str) -> dict:
    try:
        return gmail_service.users().messages().get(
            userId="me", id=message_id, format="full"
        ).execute()
    except Exception as e:
        logger.error(f"Failed to get message {message_id}: {e}")
        return {"payload": {"headers": [], "body": {}}}

def get_latest_history_id_from_messages(messages: List[dict]) -> Optional[str]:
    latest = 0
    for msg in messages:
        hid = msg.get("historyId")
        try:
            val = int(hid)
            if val > latest:
                latest = val
        except (TypeError, ValueError):
            continue
    return str(latest) if latest else None

def thread_has_changed_since(
    gmail_service, thread_id: str, start_history_id: str
) -> Optional[bool]:
    """FIX 9: pass historyTypes to reduce API response volume."""
    if not start_history_id:
        return None
    try:
        page_token = None
        while True:
            response = gmail_service.users().history().list(
                userId="me",
                startHistoryId=start_history_id,
                pageToken=page_token,
                historyTypes=["messageAdded", "messageDeleted"],   # FIX 9
                fields="history(id,messagesAdded,messagesDeleted,messages),nextPageToken",
            ).execute()
            for history in response.get("history", []):
                for event_list in ("messagesAdded", "messagesDeleted", "messages"):
                    for item in history.get(event_list, []):
                        if item.get("threadId") == thread_id:
                            return True
            page_token = response.get("nextPageToken")
            if not page_token:
                break
        return False
    except HttpError as e:
        status = getattr(getattr(e, 'resp', None), 'status', None)
        if status == 410:
            logger.warning(
                f"Gmail history expired (id={start_history_id}); "
                f"full refresh required for thread {thread_id}."
            )
            return None
        logger.error(f"History API error for thread {thread_id}: {e}")
        return None
    except Exception as e:
        logger.error(f"Failed to query Gmail history for thread {thread_id}: {e}")
        return None

def get_message_headers(message: dict) -> dict:
    headers = {}
    for h in message.get("payload", {}).get("headers", []):
        name = h["name"].lower()
        if name in ("from", "to", "cc", "subject", "date"):
            headers[name] = h["value"]
    return headers

def get_message_body(message: dict) -> str:
    def decode_text(data: str) -> str:
        return base64.urlsafe_b64decode(data).decode("utf-8", errors="replace")

    def decode_html(data: str) -> str:
        html = base64.urlsafe_b64decode(data).decode("utf-8", errors="replace")
        html = re.sub(r'<br\s*/?>', '\n', html)
        html = re.sub(r'</p>',      '\n', html)
        html = re.sub(r'</div>',    '\n', html)
        html = re.sub(r'<[^>]+>',   '',   html)
        html = re.sub(r'\n\s*\n', '\n\n', html)
        return html.strip()

    body_texts: Dict[str, List[str]] = {"plain": [], "html": []}

    def traverse_parts(part: dict):
        if not part:
            return
        mime_type = part.get("mimeType", "")
        filename  = part.get("filename")
        data      = part.get("body", {}).get("data", "")
        if mime_type == "text/plain" and not filename and data:
            body_texts["plain"].append(decode_text(data))
        elif mime_type == "text/html" and not filename and data:
            body_texts["html"].append(decode_html(data))
        for child in part.get("parts") or []:
            traverse_parts(child)

    traverse_parts(message.get("payload", {}))
    if body_texts["plain"]:
        return "\n\n".join(body_texts["plain"]).strip()
    if body_texts["html"]:
        return "\n\n".join(body_texts["html"]).strip()
    return ""

def get_attachments(gmail_service, message: dict) -> List[dict]:
    attachments: List[dict] = []
    def find_attachments(parts):
        if not parts:
            return
        for part in parts:
            if part.get("filename") and part.get("body", {}).get("attachmentId"):
                attachments.append({
                    "filename":     part["filename"],
                    "mimeType":     part.get("mimeType", ""),
                    "attachmentId": part["body"]["attachmentId"],
                    "messageId":    message["id"],
                })
            find_attachments(part.get("parts", []))
    find_attachments(message.get("payload", {}).get("parts", []))
    return attachments

def download_attachment(gmail_service, message_id: str, attachment_id: str) -> bytes:
    for attempt in range(3):
        try:
            att = gmail_service.users().messages().attachments().get(
                userId="me", messageId=message_id, id=attachment_id
            ).execute()
            return base64.urlsafe_b64decode(att["data"])
        except Exception as e:
            if attempt < 2:
                logger.warning(f"Download attempt {attempt+1} failed, retrying…")
                time.sleep(2)
            else:
                raise

# -----------------------------------------------------------------------
# TESSERACT SETUP – cached  (FIX 3)
# -----------------------------------------------------------------------
_TESSERACT_AVAILABLE: Optional[bool] = None   # None = not yet checked

def setup_tesseract() -> bool:
    global _TESSERACT_AVAILABLE
    if _TESSERACT_AVAILABLE is not None:
        return _TESSERACT_AVAILABLE

    try:
        import pytesseract
        username = os.getenv('USERNAME', '')
        tesseract_paths = [
            r'C:\Program Files\Tesseract-OCR\tesseract.exe',
            r'C:\Program Files (x86)\Tesseract-OCR\tesseract.exe',
        ]
        if username:
            tesseract_paths.append(
                rf'C:\Users\{username}\AppData\Local\Programs\Tesseract-OCR\tesseract.exe'
            )
        for path in tesseract_paths:
            if os.path.exists(path):
                pytesseract.pytesseract.tesseract_cmd = path
                logger.info(f"Tesseract found at: {path}")
                _TESSERACT_AVAILABLE = True
                return True
        try:
            subprocess.run(['tesseract', '--version'], capture_output=True, text=True, check=True)
            pytesseract.pytesseract.tesseract_cmd = 'tesseract'
            logger.info("Tesseract found in system PATH")
            _TESSERACT_AVAILABLE = True
            return True
        except (FileNotFoundError, subprocess.CalledProcessError):
            pass
        logger.warning("Tesseract not found. Image OCR disabled.")
        _TESSERACT_AVAILABLE = False
        return False
    except ImportError:
        logger.warning("pytesseract not installed.")
        _TESSERACT_AVAILABLE = False
        return False

# -----------------------------------------------------------------------
# TEXT EXTRACTION
# -----------------------------------------------------------------------
def extract_text_from_old_doc(file_bytes: bytes, file_name: str) -> str:
    with tempfile.NamedTemporaryFile(suffix='.doc', delete=False) as tmp:
        tmp.write(file_bytes)
        tmp_path = tmp.name
    try:
        result = subprocess.run(['antiword', tmp_path], capture_output=True, text=True, timeout=30)
        if result.returncode == 0 and result.stdout.strip():
            return result.stdout.strip()
    except FileNotFoundError:
        pass
    except Exception as e:
        logger.debug(f"antiword failed: {e}")
    finally:
        try:
            os.unlink(tmp_path)
        except OSError:
            pass

    try:
        import olefile
        ole = olefile.OleFileIO(io.BytesIO(file_bytes))
        parts = []
        for table_name in ['1Table', '0Table']:
            if ole.exists(table_name):
                data = ole.openstream(table_name).read()
                text = data.decode('utf-8', errors='replace')
                text = ''.join(c for c in text if c.isprintable() or c in '\n\r\t')
                if text.strip():
                    parts.append(text)
        ole.close()
        if parts:
            return '\n'.join(parts)
    except ImportError:
        logger.debug("olefile not installed")
    except Exception as e:
        logger.debug(f"olefile failed: {e}")

    try:
        text = file_bytes.decode('utf-8', errors='replace')
        text = re.sub(r'[^\x20-\x7E\n\r\t\u00A0-\u00FF]', ' ', text)
        text = re.sub(r'\s+', ' ', text).strip()
        if len(text) > 50:
            return text[:10000]
    except Exception:
        pass
    return f"[Unable to extract text from old .doc: {file_name}]"

def extract_text_from_docx(file_bytes: bytes, file_name: str) -> str:
    try:
        import docx
        doc   = docx.Document(io.BytesIO(file_bytes))
        parts = []
        for para in doc.paragraphs:
            if para.text.strip():
                parts.append(para.text.strip())
        for idx, table in enumerate(doc.tables):
            rows = [" | ".join(c.text.strip() for c in row.cells) for row in table.rows]
            rows = [r for r in rows if r.strip()]
            if rows:
                parts.append(f"\n--- Table {idx+1} ---")
                parts.extend(rows)
        for section in doc.sections:
            for para in (section.header or section.header).paragraphs if section.header else []:
                if para.text.strip():
                    parts.append(f"[Header] {para.text.strip()}")
            for para in (section.footer.paragraphs if section.footer else []):
                if para.text.strip():
                    parts.append(f"[Footer] {para.text.strip()}")
        return "\n".join(parts) if parts else "[Empty document]"
    except Exception as e:
        logger.error(f"DOCX extraction failed for {file_name}: {e}")
        if 'not a Word file' in str(e):
            return extract_text_from_old_doc(file_bytes, file_name)
        return f"[DOCX extraction failed: {e}]"

def extract_text_from_pdf(file_bytes: bytes, file_name: str) -> str:
    parts = []
    try:
        import pdfplumber
        with pdfplumber.open(io.BytesIO(file_bytes)) as pdf:
            if pdf.metadata and pdf.metadata.get('encrypted'):
                return f"[Password Protected PDF: {file_name}]"
            for page in pdf.pages:
                text = page.extract_text()
                if text and text.strip():
                    parts.append(text.strip())
    except Exception as e:
        if 'password' in str(e).lower() or 'encrypt' in str(e).lower():
            return f"[Password Protected PDF: {file_name}]"
        logger.debug(f"pdfplumber failed for {file_name}: {e}")

    if not parts:
        try:
            from PyPDF2 import PdfReader
            reader = PdfReader(io.BytesIO(file_bytes))
            if reader.is_encrypted:
                return f"[Password Protected PDF: {file_name}]"
            for page in reader.pages:
                text = page.extract_text()
                if text and text.strip():
                    parts.append(text.strip())
        except Exception as e:
            if 'password' in str(e).lower() or 'encrypt' in str(e).lower():
                return f"[Password Protected PDF: {file_name}]"
            logger.debug(f"PyPDF2 failed for {file_name}: {e}")

    if not parts and len(file_bytes) > 100_000 and setup_tesseract():
        try:
            from pdf2image import convert_from_bytes
            from pdf2image.pdf2image import pdfinfo_from_bytes
            import pytesseract
            pdfinfo_from_bytes(file_bytes)
            for image in convert_from_bytes(file_bytes, dpi=150):
                text = pytesseract.image_to_string(image.convert('L'))
                if text.strip():
                    parts.append(text.strip())
        except Exception as e:
            if "poppler" in str(e).lower():
                return f"[PDF OCR skipped – poppler not installed: {file_name}]"
            logger.error(f"PDF OCR failed for {file_name}: {e}")

    if parts:
        return "\n\n".join(parts)
    return (
        f"[Small PDF – No extractable text: {file_name}]"
        if len(file_bytes) < 5000
        else f"[PDF – No extractable text: {file_name} ({len(file_bytes)} bytes)]"
    )

def extract_text_from_image(file_bytes: bytes, file_name: str) -> str:
    if not setup_tesseract():
        return "[Image OCR not available – install Tesseract and pytesseract]"
    try:
        from PIL import Image, ImageEnhance
        import pytesseract
        image = Image.open(io.BytesIO(file_bytes))
        if image.mode == 'RGBA':
            bg    = Image.new('RGBA', image.size, (255, 255, 255))
            image = Image.alpha_composite(bg, image)
        image = image.convert('L')
        image = ImageEnhance.Contrast(image).enhance(2.0)
        image = ImageEnhance.Sharpness(image).enhance(2.0)
        results = []
        for cfg in ['--psm 6', '--psm 3', '--psm 4', '--oem 1 --psm 6']:
            try:
                text = pytesseract.image_to_string(image, config=cfg).strip()
                if text and text not in results:
                    results.append(text)
            except Exception as e:
                logger.debug(f"OCR config {cfg} failed: {e}")
        return "\n\n[OCR Results]\n\n".join(results) if results else "[No text found in image]"
    except Exception as e:
        logger.error(f"Image OCR failed for {file_name}: {e}")
        return f"[Image OCR failed: {e}]"

# -----------------------------------------------------------------------
# FIX 4: worker task uses a temp-file path, not raw bytes, to avoid
#         large pickle copies across process boundaries.
# -----------------------------------------------------------------------
def _extract_text_task(tmp_path: str, file_name: str) -> str:
    """Worker entry point. Reads bytes from tmp_path, then deletes it."""
    try:
        with open(tmp_path, 'rb') as fh:
            file_bytes = fh.read()
    finally:
        try:
            os.unlink(tmp_path)
        except OSError:
            pass
    return extract_text_locally_inner(file_bytes, file_name)

def extract_text_locally(file_bytes: bytes, file_name: str, timeout_seconds: int = 60) -> str:
    """Write bytes to a temp file, submit path to worker pool, return result."""
    suffix = os.path.splitext(file_name)[-1] or '.bin'
    try:
        with tempfile.NamedTemporaryFile(suffix=suffix, delete=False) as tmp:
            tmp.write(file_bytes)
            tmp_path = tmp.name
    except OSError as e:
        logger.error(f"Could not write temp file for {file_name}: {e}")
        return extract_text_locally_inner(file_bytes, file_name)  # fallback: in-process

    try:
        executor = get_executor()
        future   = executor.submit(_extract_text_task, tmp_path, file_name)
        return future.result(timeout=timeout_seconds)
    except concurrent.futures.TimeoutError:
        logger.warning(f"Extraction of {file_name} timed out after {timeout_seconds}s")
        try:
            os.unlink(tmp_path)
        except OSError:
            pass
        return f"[Extraction timed out after {timeout_seconds}s]"
    except Exception as e:
        logger.error(f"Extraction failed for {file_name}: {e}")
        return f"[Extraction failed: {e}]"

def extract_text_locally_inner(file_bytes: bytes, file_name: str) -> str:
    """FIX 15: all optional imports are guarded with try/except here."""
    if not file_bytes:
        return "[Empty file]"
    file_ext = file_name.rsplit('.', 1)[-1].lower() if '.' in file_name else ''
    logger.info(f"Extracting {file_name} ({file_ext}, {len(file_bytes)} bytes)")

    if file_ext == 'doc':
        return extract_text_from_old_doc(file_bytes, file_name)
    if file_ext == 'docx':
        return extract_text_from_docx(file_bytes, file_name)
    if file_ext == 'pdf':
        return extract_text_from_pdf(file_bytes, file_name)
    if file_ext in ('xlsx', 'xls', 'xlsm'):
        try:
            import openpyxl
            wb    = openpyxl.load_workbook(io.BytesIO(file_bytes), read_only=True, data_only=True)
            buf   = []
            total = 0
            MAX   = 2000
            for sheet_name in wb.sheetnames:
                ws = wb[sheet_name]
                buf.append(f"\n{'='*40}\nSheet: {sheet_name}\n{'='*40}")
                for row in ws.iter_rows(values_only=True):
                    if total >= MAX:
                        buf.append(f"[Truncated after {MAX} rows]")
                        break
                    if all(c is None for c in row):
                        continue
                    r = " | ".join(str(c) if c is not None else '' for c in row)
                    if r.strip():
                        buf.append(r)
                        total += 1
                if total >= MAX:
                    break
            wb.close()
            return "\n".join(buf).strip() or f"[Empty spreadsheet: {file_name}]"
        except ImportError:
            return "[openpyxl not installed – cannot read spreadsheet]"
        except Exception as e:
            return f"[Spreadsheet extraction failed: {e}]"
    if file_ext in ('png', 'jpg', 'jpeg', 'gif', 'bmp', 'tiff', 'tif', 'webp'):
        return extract_text_from_image(file_bytes, file_name)
    if file_ext == 'csv':
        try:
            f      = io.StringIO(file_bytes.decode('utf-8', errors='replace'))
            reader = csv.reader(f)
            rows   = [row for row in reader if any(c.strip() for c in row)]
            return "\n".join(" | ".join(row) for row in rows[:500]) if rows else "[Empty CSV]"
        except Exception as e:
            return f"[CSV parsing error: {e}]"
    if file_ext in ('txt', 'rtf', 'text', 'log', 'md', 'py', 'js', 'html', 'xml', 'json', 'cfg', 'ini'):
        for enc in ('utf-8', 'latin-1', 'cp1252', 'ascii'):
            try:
                return file_bytes.decode(enc).strip()
            except (UnicodeDecodeError, LookupError):
                continue
        return file_bytes.decode('utf-8', errors='replace').strip()
    # Unknown binary
    try:
        text = file_bytes.decode('utf-8', errors='replace').strip()
        if len(text) > 10:
            return text[:10000]
    except Exception:
        pass
    return f"[Binary file type: {file_ext} – {len(file_bytes)} bytes]"

# -----------------------------------------------------------------------
# MYSQL HELPERS
# -----------------------------------------------------------------------
def load_existing_rows(db_conn, thread_ids: Optional[Set[str]] = None) -> Dict[str, dict]:
    """FIX 7: optionally filter by a set of thread_ids to avoid loading everything."""
    cursor = db_conn.cursor(dictionary=True)
    if thread_ids:
        placeholders = ",".join(["%s"] * len(thread_ids))
        cursor.execute(
            f"SELECT id, thread_id, msg_count, message_ids, history_ids, latest_history_id "
            f"FROM threads WHERE thread_id IN ({placeholders})",
            list(thread_ids),
        )
    else:
        cursor.execute(
            "SELECT id, thread_id, msg_count, message_ids, history_ids, latest_history_id "
            "FROM threads"
        )
    rows   = cursor.fetchall()
    cursor.close()
    result = {}
    for row in rows:
        tid  = row["thread_id"]
        result[tid] = {
            "row_id":            row["id"],
            "msg_count":         row["msg_count"],
            "message_ids":       set(filter(None, (row["message_ids"]  or "").split(","))),
            "history_ids":       set(filter(None, (row.get("history_ids") or "").split(","))),
            "latest_history_id": row.get("latest_history_id"),
        }
    return result

def trigger_backend_tender_scan():
    """Trigger immediate tender matching scan on backend container when a new email is inserted."""
    backend_url = os.getenv("BACKEND_SCAN_URL", "http://backend:6003/api/sync/scan-new")
    try:
        import urllib.request
        req = urllib.request.Request(backend_url, data=b"", method="POST")
        with urllib.request.urlopen(req, timeout=5) as response:
            logger.info("Triggered immediate backend tender matching scan for new incoming email.")
    except Exception:
        try:
            import urllib.request
            req = urllib.request.Request("http://localhost:6003/api/sync/scan-new", data=b"", method="POST")
            with urllib.request.urlopen(req, timeout=5) as response:
                logger.info("Triggered immediate backend tender matching scan via localhost.")
        except Exception:
            pass

def upsert_thread(db_conn, data: dict, existing: Optional[dict] = None):
    """FIX 6: validate column names against allowlist before building SQL."""
    bad = set(data.keys()) - _ALLOWED_THREAD_COLUMNS
    if bad:
        raise ValueError(f"upsert_thread called with unknown column(s): {bad}")

    cursor = db_conn.cursor()
    is_new = existing is None
    if existing:
        set_clause = ", ".join(f"{k}=%s" for k in data)
        values     = list(data.values()) + [existing["row_id"]]
        cursor.execute(f"UPDATE threads SET {set_clause} WHERE id=%s", values)
    else:
        columns      = ", ".join(data.keys())
        placeholders = ", ".join(["%s"] * len(data))
        cursor.execute(
            f"INSERT INTO threads ({columns}) VALUES ({placeholders})",
            list(data.values()),
        )
    db_conn.commit()
    cursor.close()

    if is_new:
        trigger_backend_tender_scan()

def update_processing_status(db_conn, **kwargs):
    cursor = db_conn.cursor()
    cursor.execute(
        "SELECT id FROM processing_status WHERE status='running' ORDER BY id DESC LIMIT 1"
    )
    result = cursor.fetchone()
    if result:
        set_clause = ", ".join(f"{k}=%s" for k in kwargs)
        values     = list(kwargs.values()) + [result[0]]
        cursor.execute(f"UPDATE processing_status SET {set_clause} WHERE id=%s", values)
    else:
        kwargs.setdefault('status',     'running')
        kwargs.setdefault('start_time', datetime.now())
        columns      = ", ".join(kwargs.keys())
        placeholders = ", ".join(["%s"] * len(kwargs))
        cursor.execute(
            f"INSERT INTO processing_status ({columns}) VALUES ({placeholders})",
            list(kwargs.values()),
        )
    db_conn.commit()
    cursor.close()

def get_processing_status(db_conn) -> dict:
    cursor = db_conn.cursor(dictionary=True)
    cursor.execute("SELECT * FROM processing_status ORDER BY id DESC LIMIT 1")
    result = cursor.fetchone()
    cursor.close()
    return result or {}

def parse_email_date(date_str: str) -> Optional[str]:
    if not date_str:
        return None
    try:
        dt = parsedate_to_datetime(date_str).astimezone(timezone.utc)   # FIX 10
        return dt.strftime('%Y-%m-%d %H:%M:%S')
    except Exception as e:
        logger.warning(f"Could not parse date '{date_str}': {e}")
        return None

# -----------------------------------------------------------------------
# FIX 12: CC address extraction helper
# -----------------------------------------------------------------------
_EMAIL_RE = re.compile(r'[\w.+-]+@[\w.-]+\.\w+')

def extract_bare_emails(raw_cc: Optional[str]) -> Set[str]:
    """Return set of bare email addresses from a CC header value."""
    if not raw_cc:
        return set()
    return set(m.lower() for m in _EMAIL_RE.findall(raw_cc))

def load_all_email_mapping(db_conn) -> Dict[str, Tuple[str, str, str]]:
    """Load all_email table into a dict keyed by lowercased email.
    Returns {email: (company, category, sub_category), ...}."""
    cursor = db_conn.cursor(dictionary=True)
    cursor.execute(
        "SELECT LOWER(TRIM(NAME)) as name, COMPANY, CATEGORY, SUB_CATEGORY "
        "FROM all_email "
        "WHERE NAME IS NOT NULL AND NAME != '' "
        "AND COMPANY IS NOT NULL AND COMPANY != ''"
    )
    mapping: Dict[str, Tuple[str, str, str]] = {}
    for row in cursor.fetchall():
        email = row["name"]
        if email:
            mapping[email] = (
                row["COMPANY"] or "Outsider",
                row["CATEGORY"] or "Outsider",
                row["SUB_CATEGORY"] or "Outsider",
            )
    cursor.close()
    logger.info(f"Loaded {len(mapping)} sender mappings from all_email")
    return mapping

_INTERNAL_CATEGORY = 'internal'

def resolve_company_category(
    email_map: Dict[str, Tuple[str, str, str]],
    to_emails: Set[str],
    sender_emails: Iterable[str] = (),
) -> Tuple[str, str, str]:
    """Assign a SINGLE company + category + sub_category.

    Priority:
      1. Any To recipient whose mapping has CATEGORY='INTERNAL' (internal wins).
      2. Otherwise the first mapped To recipient (any category).
      3. Otherwise, sender fallback: an INTERNAL sender, else the first mapped
         sender (any category).
      4. Otherwise Outsider.
    """
    def _pick_internal_first(emails: Iterable[str]):
        external = None
        for email in sorted(emails):
            mapped = email_map.get(email.lower())
            if not mapped:
                continue
            if mapped[1] and mapped[1].strip().lower() == _INTERNAL_CATEGORY:
                return mapped
            if external is None:
                external = mapped
        return external

    hit = _pick_internal_first(to_emails)
    if not hit:
        hit = _pick_internal_first(sender_emails)
    if hit:
        return hit
    return "Outsider", "Outsider", "Outsider"

_HEADER_LINE_RE = re.compile(r'^(to|cc)\s*:\s*(.*)$', re.IGNORECASE)
_ANY_HEADER_RE  = re.compile(r'^[a-z-]{1,24}\s*:', re.IGNORECASE)
MAX_RECIPIENTS  = 100

def _extract_recipients_from_body(body: str) -> Tuple[Set[str], Set[str]]:
    """Parse To:/Cc: header lines from stored body text.
    Returns (to_emails, cc_emails)."""
    to_emails: Set[str] = set()
    cc_emails: Set[str] = set()
    if not body:
        return to_emails, cc_emails

    lines = body.split('\n')
    i = 0
    while i < len(lines):
        stripped = lines[i].strip()
        m = _HEADER_LINE_RE.match(stripped)
        if m:
            key   = m.group(1).lower()
            value = m.group(2)
            j = i + 1
            while j < len(lines):
                nxt = lines[j].strip()
                if not nxt or _ANY_HEADER_RE.match(nxt):
                    break
                value += ' ' + nxt
                j += 1
            emails = set(_EMAIL_RE.findall(value))
            if key == 'to':
                to_emails.update(emails)
            else:
                cc_emails.update(emails)
            i = j
        else:
            i += 1
    return to_emails, cc_emails

def backfill_recipients_from_body(db_conn, limit: Optional[int] = None) -> int:
    """Backfill to_details / cc_details by parsing To:/Cc: headers from stored body.
    Returns number of rows updated."""
    cursor = db_conn.cursor(dictionary=True)
    sql = ("SELECT id, body FROM threads "
           "WHERE (to_details IS NULL OR to_details = '' OR to_details = '[None]' "
           "OR cc_details IS NULL OR cc_details = '' OR cc_details = '[None]') ")
    if limit:
        sql += f"LIMIT {int(limit)}"
    cursor.execute(sql)
    rows = cursor.fetchall()
    cursor.close()
    logger.info(f"[Backfill body] Scanning {len(rows)} threads for To:/Cc: headers in body...")

    updated = 0
    for idx, row in enumerate(rows, 1):
        to_emails, cc_emails = _extract_recipients_from_body(row["body"])
        to_list = sorted(to_emails)[:MAX_RECIPIENTS]
        cc_list = sorted(cc_emails)[:MAX_RECIPIENTS]
        to_str = ", ".join(to_list) if to_list else None
        cc_str = ", ".join(cc_list) if cc_list else None
        if to_str or cc_str:
            up = db_conn.cursor()
            up.execute(
                "UPDATE threads SET to_details = %s, cc_details = %s WHERE id = %s",
                (to_str, cc_str, row["id"]),
            )
            db_conn.commit()
            up.close()
            updated += 1
        if idx % 100 == 0:
            logger.info(f"  [Backfill body] {idx}/{len(rows)} scanned, {updated} updated")
    logger.info(f"[Backfill body] Done: {updated} updated of {len(rows)}")
    return updated

def backfill_recipients_from_gmail(db_conn, gmail_service, limit: Optional[int] = None) -> int:
    """Backfill to_details / cc_details from Gmail for threads still missing them.
    Returns number of rows updated."""
    cursor = db_conn.cursor(dictionary=True)
    sql = ("SELECT id, thread_id FROM threads "
           "WHERE (to_details IS NULL OR to_details = '' OR to_details = '[None]' "
           "OR cc_details IS NULL OR cc_details = '' OR cc_details = '[None]') ")
    if limit:
        sql += f"LIMIT {int(limit)}"
    cursor.execute(sql)
    rows = cursor.fetchall()
    cursor.close()
    logger.info(f"[Backfill gmail] Fetching {len(rows)} threads from Gmail...")

    updated = 0
    failed  = 0
    for idx, row in enumerate(rows, 1):
        try:
            thread = gmail_service.users().threads().get(
                userId="me", id=row["thread_id"], format="full"
            ).execute()
            to_emails: Set[str] = set()
            cc_emails: Set[str] = set()
            for msg in thread.get("messages", []):
                headers = get_message_headers(msg)
                to_emails.update(extract_bare_emails(headers.get("to", "")))
                cc_emails.update(extract_bare_emails(headers.get("cc", "")))
            to_list = sorted(to_emails)[:MAX_RECIPIENTS]
            cc_list = sorted(cc_emails)[:MAX_RECIPIENTS]
            to_str = ", ".join(to_list) if to_list else None
            cc_str = ", ".join(cc_list) if cc_list else None
            if to_str or cc_str:
                up = db_conn.cursor()
                up.execute(
                    "UPDATE threads SET to_details = %s, cc_details = %s WHERE id = %s",
                    (to_str, cc_str, row["id"]),
                )
                db_conn.commit()
                up.close()
                updated += 1
        except Exception as e:
            failed += 1
            logger.error(f"  [Backfill gmail] Failed thread {row['thread_id']}: {e}")
        if idx % 25 == 0:
            logger.info(f"  [Backfill gmail] {idx}/{len(rows)} processed, {updated} updated, {failed} failed")
    logger.info(f"[Backfill gmail] Done: {updated} updated, {failed} failed of {len(rows)}")
    return updated

def reassign_company_category(db_conn, limit: Optional[int] = None, only_outsider: bool = False) -> int:
    """Re-run single-company assignment for existing threads using stored recipients,
    falling back to the sender when no recipient matches.
    only_outsider=True limits the run to threads currently marked 'Outsider'.
    Returns number of rows updated."""
    email_map = load_all_email_mapping(db_conn)
    cursor = db_conn.cursor(dictionary=True)
    if only_outsider:
        sql = ("SELECT id, sender, to_details FROM threads "
               "WHERE company = 'Outsider' ")
    else:
        sql = ("SELECT id, sender, to_details FROM threads "
               "WHERE to_details IS NOT NULL AND to_details != '' AND to_details != '[None]' ")
    if limit:
        sql += f"LIMIT {int(limit)}"
    cursor.execute(sql)
    rows = cursor.fetchall()
    cursor.close()
    logger.info(f"[Reassign company] Processing {len(rows)} threads ({'Outsider only' if only_outsider else 'with recipients'})...")

    updated = 0
    for idx, row in enumerate(rows, 1):
        to_emails = extract_bare_emails(row["to_details"])
        sender_emails = _EMAIL_RE.findall(row["sender"] or "")
        company, category, sub_category = resolve_company_category(
            email_map, to_emails, sender_emails
        )
        up = db_conn.cursor()
        up.execute(
            "UPDATE threads SET company=%s, category=%s, sub_category=%s WHERE id=%s",
            (company, category, sub_category, row["id"]),
        )
        db_conn.commit()
        up.close()
        updated += 1
        if idx % 500 == 0:
            logger.info(f"  [Reassign company] {idx}/{len(rows)} processed")
    logger.info(f"[Reassign company] Done: {updated} updated of {len(rows)}")
    return updated

def backfill_and_reassign_company(db_conn, gmail_service, limit: Optional[int] = None) -> dict:
    """Gmail-first recipient backfill, then body fallback, then re-assign company/category.
    Returns summary counts."""
    gmail_updated = backfill_recipients_from_gmail(db_conn, gmail_service, limit=limit)
    body_updated  = backfill_recipients_from_body(db_conn, limit=limit)
    assign_updated = reassign_company_category(db_conn, limit=None)
    summary = {
        'recipients_gmail': gmail_updated,
        'recipients_body':  body_updated,
        'company_reassigned': assign_updated,
    }
    logger.info(f"[Backfill+Reassign] Summary: {summary}")
    return summary

# -----------------------------------------------------------------------
# MAIN PROCESS
# -----------------------------------------------------------------------
def process_gmail_batch(
    gmail_service, drive_service,
    all_threads, email_map,
    start_index: int = 0, batch_size: int = 500
):
    db_conn = get_db_connection()
    logger.info(f"Starting batch from index {start_index}, size {batch_size}")
    logger.info(f"Search query: {SEARCH_QUERY}")

    total_threads = len(all_threads)
    end_index     = min(start_index + batch_size, total_threads)
    batch_threads = all_threads[start_index:end_index]
    logger.info(f"Processing batch: {start_index+1}–{end_index} of {total_threads}")

    # FIX 7: load only the rows we might touch in this batch
    batch_thread_ids = {t["id"] for t in batch_threads}
    existing_rows    = load_existing_rows(db_conn, thread_ids=batch_thread_ids)

    root_folder_id = DRIVE_FOLDER_ID or create_drive_folder(
        drive_service, f"Gmail_Attachments_{datetime.now().strftime('%Y%m%d_%H%M%S')}"
    )

    start_time     = time.time() * 1000
    processed_count = important_count = failed_count = skipped_count = 0

    for i, thread_summary in enumerate(batch_threads):
        thread_index = start_index + i + 1

        if time.time() * 1000 - start_time > MAX_RUN_TIME_MS:
            logger.warning("Time limit reached. Saving progress…")
            break

        thread_id = thread_summary["id"]
        logger.info(f"-> Thread {thread_index}/{total_threads} ({thread_id})")
        sys.stdout.flush()
        try:
            existing = existing_rows.get(thread_id)

            # Fast skip via Gmail History API
            if existing and existing.get("latest_history_id"):
                changed = thread_has_changed_since(
                    gmail_service, thread_id, existing["latest_history_id"]
                )
                if changed is False:
                    skipped_count += 1
                    logger.info(f"Skip {thread_id}: no changes.")
                    continue

            messages = get_thread_messages(gmail_service, thread_id)
            if not messages:
                logger.warning(f"No messages in thread {thread_id}")
                continue

            msg_datas = [get_message_data(gmail_service, m["id"]) for m in messages]

            current_msg_ids     = {md["id"] for md in msg_datas}
            current_history_ids = set()
            for md in msg_datas:
                hid = md.get("historyId")
                if hid:
                    current_history_ids.add(str(hid))
                else:
                    snippet          = md.get("snippet", "") or ""
                    headers_concat   = "".join(
                        h.get("value", "") for h in md.get("payload", {}).get("headers", [])
                    )
                    synthetic        = f"{md.get('id','')}//{snippet}//{headers_concat}"
                    current_history_ids.add(hashlib.sha256(synthetic.encode()).hexdigest())

            current_latest_history_id = get_latest_history_id_from_messages(msg_datas)

            # Fallback skip checks (no History API id stored yet)
            if existing and not existing.get("latest_history_id"):
                if existing.get("history_ids") and current_history_ids == existing["history_ids"]:
                    skipped_count += 1
                    continue
                if (not existing.get("history_ids")
                        and existing.get("message_ids")
                        and current_msg_ids == existing["message_ids"]):
                    skipped_count += 1
                    continue

            is_update = bool(existing)
            logger.info(
                f"{'[Update]' if is_update else '[Insert]'} "
                f"Thread {thread_id} ({thread_index}/{total_threads})"
            )

            latest_msg = msg_datas[-1]
            headers    = get_message_headers(latest_msg)
            date_str   = headers.get("date", "")
            from_field = headers.get("from", "")
            subject    = headers.get("subject", "")

            from_field_lower = (from_field or "").lower()
            if "automation@app.smartsheet.com" in from_field_lower or "tendertiger.com" in from_field_lower:
                logger.info(f"  [Skipped Blacklisted Sender] {from_field} - Subject: {subject[:60]}")
                continue

            cc_emails: Set[str] = set()   # FIX 12: bare email dedup
            to_emails: Set[str] = set()
            body_parts: List[str] = []
            all_attachment_names: List[str] = []

            for idx, msg_data in enumerate(msg_datas, 1):
                mh = get_message_headers(msg_data)
                body_parts.append(f"--- Message {idx} From: {mh.get('from','')} ---")
                body_parts.append(get_message_body(msg_data))
                cc = mh.get("cc", "")
                if cc:
                    cc_emails.update(extract_bare_emails(cc))   # FIX 12
                to = mh.get("to", "")
                if to:
                    to_emails.update(extract_bare_emails(to))
                for att in get_attachments(gmail_service, msg_data):
                    all_attachment_names.append(att["filename"])

            body_content = "\n\n".join(body_parts)
            cc_details   = ", ".join(sorted(cc_emails)) or "[None]"
            to_details   = ", ".join(sorted(to_emails)) or "[None]"

            is_important, importance_reasons = determine_importance(
                subject, body_content, from_field, all_attachment_names
            )

            # Company lookup: To recipients first, then sender fallback, one company per mail
            company, category, sub_category = resolve_company_category(
                email_map, to_emails, _EMAIL_RE.findall(from_field or "")
            )

            if is_important:
                important_count += 1
                logger.info(f"  [IMPORTANT] {subject[:100]} – {category}")
                if importance_reasons:
                    logger.info(f"  Reasons: {', '.join(importance_reasons[:3])}")

            thread_folder_id = get_or_create_thread_folder(drive_service, thread_id, root_folder_id)

            attachment_drive_link_cache: Dict[str, Optional[str]] = {}
            attachment_text_cache: Dict[str, str]                 = {}
            attachment_name_by_hash: Dict[str, str]               = {}
            processed_hashes: List[str]                            = []
            attachment_sections: List[str]                         = []

            for msg_data in msg_datas:
                for att in get_attachments(gmail_service, msg_data):
                    fname = att["filename"]
                    try:
                        att_bytes = download_attachment(
                            gmail_service, att["messageId"], att["attachmentId"]
                        )
                        att_hash = hashlib.sha256(att_bytes).hexdigest()
                    except Exception as e:
                        logger.error(f"  Download error for {fname}: {e}")
                        attachment_sections.append(f"=== {fname} ===\n[Download error: {e}]")
                        continue

                    if att_hash in attachment_drive_link_cache:
                        drive_link = attachment_drive_link_cache[att_hash]
                        ocr_text   = attachment_text_cache.get(att_hash, "")
                        attachment_sections.append(
                            f"\n{'='*60}\nATTACHMENT: {fname} (duplicate – reusing)\n"
                            f"DRIVE LINK: {drive_link or 'Upload failed'}\n{'='*60}\n"
                            f"{ocr_text or '[No extractable text]'}\n"
                        )
                        continue

                    processed_hashes.append(att_hash)
                    attachment_name_by_hash[att_hash] = fname
                    try:
                        logger.info(f"  Processing: {fname}")
                        drive_link = upload_to_drive_dedup(
                            drive_service, att_bytes, fname, thread_folder_id
                        )
                        doc_text   = extract_text_locally(att_bytes, fname)
                        attachment_drive_link_cache[att_hash] = drive_link
                        attachment_text_cache[att_hash]       = doc_text
                        attachment_sections.append(
                            f"\n{'='*60}\nATTACHMENT: {fname}\n"
                            f"DRIVE LINK: {drive_link or 'Upload failed'}\n"
                            f"FILE SIZE:  {len(att_bytes)} bytes\n{'='*60}\n"
                            f"{doc_text or '[No extractable text]'}\n"
                        )
                    except Exception as e:
                        logger.error(f"  Error processing {fname}: {e}")
                        attachment_drive_link_cache[att_hash] = None
                        attachment_text_cache[att_hash]       = f"[Error: {e}]"
                        attachment_sections.append(f"=== {fname} ===\n[Error: {e}]")

            ocr_text = "\n\n".join(attachment_sections) if attachment_sections else "[No Attachments]"
            if len(ocr_text) > OCR_MAX_CHARS:
                ocr_text = ocr_text[:OCR_MAX_CHARS] + f"\n[Truncated – total {len(ocr_text)} chars]"

            file_names  = [attachment_name_by_hash[h] for h in processed_hashes]
            drive_links = [attachment_drive_link_cache.get(h) or "" for h in processed_hashes]

            row_data = {
                "thread_id":          thread_id,
                "related_ids":        "[None]",
                "msg_count":          len(msg_datas),
                "date":               parse_email_date(date_str),
                "sender":             from_field,
                "sender_details":     from_field,
                "cc_details":         cc_details,
                "to_details":         to_details,
                "subject":            subject,
                "body":               body_content[:BODY_MAX_CHARS],
                "attach_names":       ", ".join(file_names)  or "[No Attachments]",
                "attach_links":       ", ".join(drive_links) or "[No Links]",
                "ocr_text":           ocr_text,
                "ai_summary":         "[Not analyzed]",
                "category":           category,
                "sub_category":       sub_category,
                "company":            company,
                "priority":           "high" if is_important else "medium",
                "is_important":       is_important,
                "importance_reasons": ", ".join(importance_reasons),
                "contacts":           "[Not analyzed]",
                "footprint":          "[Not analyzed]",
                "message_ids":        ",".join(current_msg_ids),
                "history_ids":        ",".join(current_history_ids),
                "latest_history_id":  current_latest_history_id,
                "drive_folder_id":    thread_folder_id,
            }

            upsert_thread(db_conn, row_data, existing_rows.get(thread_id))
            processed_count += 1

            update_processing_status(
                db_conn,
                total_threads=total_threads,
                processed_threads=thread_index,
                important_threads=important_count,
                last_thread_id=thread_id,
                last_processed_date=parse_email_date(date_str),
                batch_number=(start_index // batch_size) + 1,
            )
            logger.info(
                f"  [OK] {processed_count}/{len(batch_threads)} "
                f"(important={important_count}, skipped={skipped_count})"
            )

        except Exception as e:
            failed_count += 1
            logger.error(f"[FAIL] Thread {thread_id}: {e}", exc_info=True)

    update_processing_status(
        db_conn,
        total_threads=total_threads,
        processed_threads=end_index,
        important_threads=important_count,
        end_time=datetime.now(),
        status='completed' if end_index >= total_threads else 'partial',
    )
    db_conn.close()

    logger.info(f"\n{'='*60}")
    logger.info(f"Batch summary: processed={processed_count}, important={important_count}, "
                f"failed={failed_count}, skipped={skipped_count}")
    logger.info(f"{'='*60}\n")

    return {
        'total':             total_threads,
        'processed':         end_index,
        'batch_processed':   processed_count,
        'batch_important':   important_count,
        'batch_failed':      failed_count,
        'batch_skipped':     skipped_count,
        'continue':          end_index < total_threads,
        'last_thread_id':    batch_threads[-1]["id"] if batch_threads else None,
    }

def load_processed_thread_ids():
    try:
        conn = get_db_connection()
        cursor = conn.cursor()
        cursor.execute("SELECT DISTINCT thread_id FROM threads WHERE thread_id IS NOT NULL AND thread_id != ''")
        ids = set(row[0] for row in cursor.fetchall())
        conn.close()
        return ids
    except Exception as e:
        logger.warning(f"Could not load processed thread_ids from DB: {e}")
        return set()

# -----------------------------------------------------------------------
# CONTINUOUS RUN  (Filter already processed thread_ids from DB)
# -----------------------------------------------------------------------
def run_continuous():
    logger.info("="*70)
    logger.info("Starting enhanced continuous Gmail listener pipeline")
    logger.info("="*70)

    init_db()
    gmail_service, drive_service = get_google_services()

    processed_thread_ids = load_processed_thread_ids()
    logger.info(f"Loaded {len(processed_thread_ids)} existing processed threads from database.")

    batch_number = 1
    total_important = 0

    while True:
        try:
            # Re-query DB for any newly added thread_ids
            latest_db_ids = load_processed_thread_ids()
            processed_thread_ids.update(latest_db_ids)

            setup_conn = get_db_connection()
            email_map = load_all_email_mapping(setup_conn)
            setup_conn.close()

            # Search Gmail API for threads matching search query
            all_threads = search_threads(gmail_service, SEARCH_QUERY)
            
            # Filter ONLY threads that are not yet processed in DB
            unprocessed_threads = [t for t in all_threads if t["id"] not in processed_thread_ids]

            if not unprocessed_threads:
                logger.info(f"All {len(all_threads)} threads up-to-date. Listening for new incoming emails...")
                time.sleep(15)
                continue

            logger.info(f"\n{'#'*70}\nBATCH #{batch_number}: Processing {len(unprocessed_threads)} new incoming threads (out of {len(all_threads)} total)\n{'#'*70}\n")

            result = process_gmail_batch(
                gmail_service, drive_service,
                unprocessed_threads, email_map,
                start_index=0, batch_size=BATCH_SIZE
            )
            total_important += result.get('batch_important', 0)
            logger.info(f"Cumulative important: {total_important}")

            # Mark all processed threads in set
            for t in unprocessed_threads:
                processed_thread_ids.add(t["id"])

            batch_number += 1
            logger.info("Checking for new incoming emails in 15s...")
            time.sleep(15)

        except KeyboardInterrupt:
            logger.info(f"\nInterrupted. Total important so far: {total_important}")
            break
        except Exception as e:
            logger.error(f"Batch error: {e}", exc_info=True)
            logger.info("Waiting 30s before retry...")
            time.sleep(30)

# -----------------------------------------------------------------------
# REPORTING  (FIX 14: returns data cleanly; caller prints)
# -----------------------------------------------------------------------
def generate_report(db_conn) -> dict:
    cursor = db_conn.cursor(dictionary=True)
    cursor.execute("""
        SELECT category,
               COUNT(*) AS count,
               SUM(CASE WHEN is_important=1 THEN 1 ELSE 0 END) AS important_count
        FROM threads
        WHERE category != 'General'
        GROUP BY category
        ORDER BY count DESC
    """)
    categories = cursor.fetchall()
    cursor.execute("""
        SELECT subject, sender, date, category, importance_reasons
        FROM threads
        WHERE is_important=1
        ORDER BY date DESC
        LIMIT 20
    """)
    important_emails = cursor.fetchall()
    cursor.close()
    return {'categories': categories, 'important_emails': important_emails}

# -----------------------------------------------------------------------
# ENTRY POINT
# -----------------------------------------------------------------------
if __name__ == "__main__":
    import argparse
    parser = argparse.ArgumentParser(description='Gmail OCR Pipeline')
    parser.add_argument('--mode', choices=['single', 'continuous', 'report', 'backfill-recipients',
                                           'reassign-company'],
                        default='continuous')
    parser.add_argument('--continuous', action='store_true', help='Alias for --mode continuous')
    parser.add_argument('--batch-size',   type=int, default=500)
    parser.add_argument('--search-query', type=str, default=None)
    parser.add_argument('--source', choices=['body', 'gmail', 'all'], default='all',
                        help='backfill-recipients source: body (from stored text) or gmail (re-fetch)')
    parser.add_argument('--limit', type=int, default=None,
                        help='max rows to process for backfill-recipients')
    args = parser.parse_args()

    if args.continuous:
        args.mode = 'continuous'

    BATCH_SIZE = args.batch_size
    if args.search_query:
        SEARCH_QUERY = args.search_query
        logger.info(f"Custom search query: {SEARCH_QUERY[:200]}")

    if args.mode == 'report':
        db_conn = get_db_connection()
        try:
            report = generate_report(db_conn)    # FIX 14: data fetched before close
        finally:
            db_conn.close()
        print("\n" + "="*70)
        print("EMAIL CATEGORY SUMMARY")
        print("="*70)
        for cat in report['categories']:
            print(f"  {cat['category']}: {cat['count']} ({cat['important_count']} important)")
        print("\n" + "="*70)
        print("RECENT IMPORTANT EMAILS")
        print("="*70)
        for email in report['important_emails']:
            print(f"  [{email['date']}] {email['subject'][:80]}")
            print(f"    From: {email['sender'][:60]}")
            print(f"    Category: {email['category']}")
            print(f"    Reasons: {(email['importance_reasons'] or '')[:100]}")
            print()
    elif args.mode == 'continuous':
        run_continuous()
    elif args.mode == 'backfill-recipients':
        db_conn = get_db_connection()
        try:
            if args.source in ('body', 'all'):
                backfill_recipients_from_body(db_conn, limit=args.limit)
            if args.source in ('gmail', 'all'):
                gmail_service, _ = get_google_services()
                backfill_recipients_from_gmail(db_conn, gmail_service, limit=args.limit)
        finally:
            db_conn.close()
    elif args.mode == 'reassign-company':
        db_conn = get_db_connection()
        try:
            gmail_service, _ = get_google_services()
            backfill_and_reassign_company(db_conn, gmail_service, limit=args.limit)
        finally:
            db_conn.close()
    else:
        init_db()
        db_conn = get_db_connection()
        gmail_service, drive_service = get_google_services()
        email_map = load_all_email_mapping(db_conn)
        all_threads = search_threads(gmail_service, SEARCH_QUERY)
        result = process_gmail_batch(
            gmail_service, drive_service,
            all_threads, email_map,
            start_index=0, batch_size=BATCH_SIZE
        )
        logger.info(f"Single batch done: processed={result['batch_processed']}, "
                    f"important={result['batch_important']}")